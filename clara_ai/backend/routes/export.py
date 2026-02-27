"""Export endpoint — POST /export_report → downloadable DOCX or PDF."""

from __future__ import annotations

import io
import logging
import re
from datetime import datetime, timezone
from typing import Literal

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

logger = logging.getLogger(__name__)
router = APIRouter(tags=["export"])


# ── Request schema ────────────────────────────────────────────────────────────

class ExportRequest(BaseModel):
    report_text: str
    format: Literal["pdf", "docx"] = "docx"


# ── Helpers ───────────────────────────────────────────────────────────────────

def _strip_md(text: str) -> str:
    """Remove Markdown syntax, returning plain text."""
    text = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"^\|.*\|$", "", text, flags=re.MULTILINE)
    text = re.sub(r"^[-|]+$", "", text, flags=re.MULTILINE)
    text = re.sub(r"^>\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ── DOCX builder ──────────────────────────────────────────────────────────────

def _build_docx(report_text: str) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as exc:
        raise RuntimeError("python-docx not installed") from exc

    doc = Document()

    # Title styling
    section = doc.sections[0]
    section.page_width = int(8.5 * 914400)
    section.page_height = int(11 * 914400)

    for para in doc.paragraphs:
        para.clear()

    lines = report_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        # H1
        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1

        # H2
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1

        # H3
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1

        # Horizontal rule
        elif line.strip() == "---":
            doc.add_paragraph("─" * 60)
            i += 1

        # Blockquote
        elif line.startswith("> "):
            p = doc.add_paragraph(style="Quote")
            p.add_run(line[2:].strip())
            i += 1

        # Table (collect consecutive | lines)
        elif line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                row = lines[i]
                if not re.match(r"^\|[-| ]+\|$", row):
                    table_lines.append(row)
                i += 1
            if table_lines:
                rows = [
                    [cell.strip() for cell in r.strip("|").split("|")]
                    for r in table_lines
                ]
                col_count = max(len(r) for r in rows)
                tbl = doc.add_table(rows=len(rows), cols=col_count)
                tbl.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        if ci < col_count:
                            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", cell)
                            clean = re.sub(r"\*(.+?)\*", r"\1", clean)
                            tbl.cell(ri, ci).text = clean
                            if ri == 0:
                                run = tbl.cell(ri, ci).paragraphs[0].runs
                                if run:
                                    run[0].bold = True

        # Numbered list
        elif re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            text = re.sub(r"^\d+\.\s*", "", line)
            _add_rich_run(p, text)
            i += 1

        # Bullet list
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_rich_run(p, line[2:].strip())
            i += 1

        # Empty line
        elif line.strip() == "":
            doc.add_paragraph("")
            i += 1

        # Normal paragraph
        else:
            p = doc.add_paragraph()
            _add_rich_run(p, line)
            i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_rich_run(para, text: str) -> None:
    """Parse **bold** and *italic* inline and add styled runs."""
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            para.add_run(text[pos:m.start()])
        if m.group(0).startswith("**"):
            para.add_run(m.group(2)).bold = True
        elif m.group(0).startswith("*"):
            para.add_run(m.group(3)).italic = True
        else:
            r = para.add_run(m.group(4))
            r.font.name = "Courier New"
        pos = m.end()
    if pos < len(text):
        para.add_run(text[pos:])


# ── PDF builder ───────────────────────────────────────────────────────────────

def _build_pdf(report_text: str) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer,
            Table, TableStyle, HRFlowable,
        )
    except ImportError as exc:
        raise RuntimeError("reportlab not installed") from exc

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=22 * mm,
        bottomMargin=22 * mm,
    )

    styles = getSampleStyleSheet()
    style_h1 = ParagraphStyle(
        "H1", parent=styles["Heading1"],
        fontSize=16, leading=22, textColor=colors.HexColor("#1a237e"),
        spaceAfter=8,
    )
    style_h2 = ParagraphStyle(
        "H2", parent=styles["Heading2"],
        fontSize=13, leading=18, textColor=colors.HexColor("#283593"),
        spaceBefore=14, spaceAfter=6,
    )
    style_h3 = ParagraphStyle(
        "H3", parent=styles["Heading3"],
        fontSize=11, leading=15, textColor=colors.HexColor("#3949ab"),
        spaceBefore=8, spaceAfter=4,
    )
    style_body = ParagraphStyle(
        "Body", parent=styles["Normal"],
        fontSize=9.5, leading=14, textColor=colors.HexColor("#212121"),
        spaceAfter=4,
    )
    style_quote = ParagraphStyle(
        "Quote", parent=styles["Normal"],
        fontSize=9.5, leading=14,
        leftIndent=18, textColor=colors.HexColor("#37474f"),
        borderPad=4, borderColor=colors.HexColor("#7986cb"),
        borderWidth=2, borderRadius=3,
        spaceAfter=6,
    )
    style_code = ParagraphStyle(
        "Code", parent=styles["Code"],
        fontSize=8.5, leading=13, textColor=colors.HexColor("#546e7a"),
        spaceAfter=4,
    )
    style_bullet = ParagraphStyle(
        "Bullet", parent=styles["Normal"],
        fontSize=9.5, leading=14, leftIndent=14, firstLineIndent=-10,
        spaceAfter=3,
    )

    flowables = []
    lines = report_text.splitlines()
    i = 0

    def _escape(t: str) -> str:
        t = t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        t = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", t)
        t = re.sub(r"\*(.+?)\*", r"<i>\1</i>", t)
        t = re.sub(r"`(.+?)`", r'<font name="Courier">\1</font>', t)
        return t

    while i < len(lines):
        line = lines[i]

        if line.startswith("# "):
            flowables.append(Paragraph(_escape(line[2:].strip()), style_h1))
            i += 1

        elif line.startswith("## "):
            flowables.append(Paragraph(_escape(line[3:].strip()), style_h2))
            i += 1

        elif line.startswith("### "):
            flowables.append(Paragraph(_escape(line[4:].strip()), style_h3))
            i += 1

        elif line.strip() == "---":
            flowables.append(HRFlowable(width="100%", thickness=0.6,
                                         color=colors.HexColor("#9fa8da"), spaceAfter=6))
            i += 1

        elif line.startswith("> "):
            flowables.append(Paragraph(_escape(line[2:].strip()), style_quote))
            i += 1

        elif line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                row = lines[i]
                if not re.match(r"^\|[-| ]+\|$", row):
                    table_lines.append(row)
                i += 1
            if table_lines:
                rows = [
                    [Paragraph(_escape(c.strip()), style_body)
                     for c in r.strip("|").split("|")]
                    for r in table_lines
                ]
                col_count = max(len(r) for r in rows)
                rows = [r + [Paragraph("", style_body)] * (col_count - len(r)) for r in rows]
                tbl = Table(rows, hAlign="LEFT")
                tbl.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8eaf6")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1a237e")),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#c5cae9")),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f3f4ff")]),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ]))
                flowables.append(tbl)
                flowables.append(Spacer(1, 6))

        elif re.match(r"^\d+\.\s", line):
            text = re.sub(r"^\d+\.\s*", "", line)
            flowables.append(Paragraph(f"• {_escape(text)}", style_bullet))
            i += 1

        elif line.startswith("- ") or line.startswith("* "):
            flowables.append(Paragraph(f"• {_escape(line[2:].strip())}", style_bullet))
            i += 1

        elif line.strip() == "":
            flowables.append(Spacer(1, 5))
            i += 1

        elif line.startswith("*This dossier"):
            flowables.append(Paragraph(_escape(line), style_code))
            i += 1

        else:
            flowables.append(Paragraph(_escape(line), style_body))
            i += 1

    doc.build(flowables)
    return buf.getvalue()


# ── Endpoint ──────────────────────────────────────────────────────────────────

@router.post("/export_report")
async def export_report(req: ExportRequest) -> StreamingResponse:
    """Convert report Markdown to DOCX or PDF and return as download."""
    ts = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    fmt = req.format.lower()

    try:
        if fmt == "docx":
            data = _build_docx(req.report_text)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = f"clara_incident_report_{ts}.docx"

        elif fmt == "pdf":
            data = _build_pdf(req.report_text)
            media_type = "application/pdf"
            filename = f"clara_incident_report_{ts}.pdf"

        else:
            raise HTTPException(status_code=400, detail="format must be 'pdf' or 'docx'")

    except RuntimeError as exc:
        logger.error("Export failed: %s", exc)
        raise HTTPException(status_code=500, detail=str(exc)) from exc

    return StreamingResponse(
        io.BytesIO(data),
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
